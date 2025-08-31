# can you add html, css code to this fastapi code ,use boostrap css and tailwind css make it visually good
import asyncio
import pandas as pd
import numpy as np
from typing import List, Dict, Any
from autogen_agentchat.agents import AssistantAgent, UserProxyAgent
from autogen_agentchat.teams import RoundRobinGroupChat
from autogen_agentchat.conditions import TextMentionTermination
from autogen_ext.models.openai import OpenAIChatCompletionClient
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt
import seaborn as sns
from datetime import datetime
import os
from dotenv import load_dotenv

load_dotenv()

import os, sys, subprocess

def launch_streamlit():
    if os.environ.get("COLAB", "0") == "1":
        print("Skipping local Streamlit launch in Colab.")
        return
    try:
        subprocess.run([
            sys.executable, "-m", "streamlit", "run", "streamlit_app.py",
            "--server.port", "8501", "--server.address", "localhost",
            "--browser.gatherUsageStats", "false"
        ])
    except Exception as e:
        print("Error launching Streamlit:", e)

class DataLoaderAgent(AssistantAgent):
    """Agent responsible for loading and validating the dataset."""
    
    def __init__(self, model_client):
        super().__init__(
            name="DataLoader",
            model_client=model_client,
            system_message="""You are the DataLoader Agent. Your role is to:
            1. Load datasets from specified file paths
            2. Validate data integrity and format
            3. Provide basic dataset information
            4. Handle any data loading errors gracefully
            
            Always respond with clear, actionable information about the dataset.
            
            When asked to load data, use the load_data function to:
            - Check if the file exists
            - Load the CSV data using pandas
            - Validate the data structure
            - Report any issues found
            
            Provide comprehensive dataset overview including:
            - Number of rows and columns
            - Column names and data types
            - Missing value counts
            - Memory usage
            - Sample data preview"""
        )
    
    def load_data(self, file_path: str) -> pd.DataFrame:
        """Load data from the specified file path."""
        try:
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"File not found at {file_path}")
            
            # Load data
            data = pd.read_csv(file_path)
            
            # Handle duplicate column names
            if len(data.columns) != len(set(data.columns)):
                print("⚠️ Warning: Duplicate column names detected. Renaming duplicates...")
                # Get duplicate columns
                duplicate_cols = data.columns[data.columns.duplicated()].tolist()
                print(f"   Duplicate columns: {duplicate_cols}")
                
                # Rename duplicate columns by adding suffix
                new_columns = []
                seen_columns = {}
                
                for col in data.columns:
                    if col in seen_columns:
                        seen_columns[col] += 1
                        new_columns.append(f"{col}.{seen_columns[col]}")
                    else:
                        seen_columns[col] = 0
                        new_columns.append(col)
                
                data.columns = new_columns
                print(f"   Columns renamed to: {list(data.columns)}")
            
            # Clean empty rows (rows with all NaN values)
            initial_rows = len(data)
            data = data.dropna(how='all')
            if len(data) < initial_rows:
                print(f"   Removed {initial_rows - len(data)} completely empty rows")
            
            # Clean empty columns (columns with all NaN values)
            initial_cols = len(data.columns)
            data = data.dropna(axis=1, how='all')
            if len(data.columns) < initial_cols:
                print(f"   Removed {initial_cols - len(data.columns)} completely empty columns")
            
            return data
            
        except Exception as e:
            raise Exception(f"Error loading data: {e}")
    
    def get_data_info(self, data: pd.DataFrame) -> Dict[str, Any]:
        """Get comprehensive information about the dataset."""
        info = {
            "shape": data.shape,
            "columns": list(data.columns),
            "data_types": dict(data.dtypes),
            "missing_values": dict(data.isnull().sum()),
            "memory_usage": data.memory_usage(deep=True).sum() / 1024,  # KB
            "sample_data": data.head(3).to_dict()
        }
        return info

class DataAnalyzerAgent(AssistantAgent):
    """Agent responsible for statistical analysis and data insights."""
    
    def __init__(self, model_client):
        super().__init__(
            name="DataAnalyzer",
            model_client=model_client,
            system_message="""You are the DataAnalyzer Agent. Your role is to:
            1. Perform comprehensive statistical analysis
            2. Generate descriptive statistics for all variables
            3. Identify patterns, trends, and anomalies
            4. Provide data-driven insights and observations
            
            When analyzing data, focus on:
            - Numerical variables: mean, median, std, quartiles, correlations
            - Categorical variables: frequency counts, distributions, unique values
            - Missing data patterns and data quality assessment
            - Outlier detection and analysis
            - Relationship analysis between variables
            
            Always provide clear, interpretable statistical information with business context.
            Use the analyze_data function to perform comprehensive analysis."""
        )
    
    def analyze_data(self, data: pd.DataFrame) -> Dict[str, Any]:
        """Perform comprehensive data analysis."""
        analysis = {}
        
        try:
            # Basic statistics
            analysis["shape"] = data.shape
            analysis["info"] = data.info()
            
            # Descriptive statistics for numerical columns
            numerical_cols = data.select_dtypes(include=[np.number]).columns
            if len(numerical_cols) > 0:
                # Handle potential errors in describe()
                try:
                    desc_stats = data[numerical_cols].describe()
                    analysis["numerical_stats"] = desc_stats.to_dict()
                except Exception as e:
                    print(f"Warning: Could not generate numerical stats: {e}")
                    analysis["numerical_stats"] = {}
                
                # Handle potential errors in correlation
                try:
                    if len(numerical_cols) > 1:
                        corr_matrix = data[numerical_cols].corr()
                        analysis["correlation_matrix"] = corr_matrix.to_dict()
                    else:
                        analysis["correlation_matrix"] = {}
                except Exception as e:
                    print(f"Warning: Could not generate correlation matrix: {e}")
                    analysis["correlation_matrix"] = {}
            
            # Categorical analysis
            categorical_cols = data.select_dtypes(include=['object']).columns
            if len(categorical_cols) > 0:
                analysis["categorical_stats"] = {}
                for col in categorical_cols:
                    try:
                        value_counts = data[col].value_counts()
                        analysis["categorical_stats"][col] = {
                            "unique_values": data[col].nunique(),
                            "value_counts": value_counts.head(10).to_dict() if len(value_counts) > 0 else {},
                            "missing_count": data[col].isnull().sum()
                        }
                    except Exception as e:
                        print(f"Warning: Could not analyze categorical column {col}: {e}")
                        analysis["categorical_stats"][col] = {
                            "unique_values": 0,
                            "value_counts": {},
                            "missing_count": data[col].isnull().sum()
                        }
            
            # Missing data analysis
            analysis["missing_data"] = {
                "total_missing": data.isnull().sum().sum(),
                "missing_by_column": dict(data.isnull().sum()),
                "missing_percentage": dict((data.isnull().sum() / len(data)) * 100)
            }
            
        except Exception as e:
            print(f"Error in data analysis: {e}")
            analysis = {
                "shape": data.shape,
                "error": str(e),
                "missing_data": {
                    "total_missing": data.isnull().sum().sum(),
                    "missing_by_column": dict(data.isnull().sum())
                }
            }
        
        return analysis

class DataVisualizerAgent(AssistantAgent):
    """Agent responsible for creating data visualizations."""
    
    def __init__(self, model_client):
        super().__init__(
            name="DataVisualizer",
            model_client=model_client,
            system_message="""You are the DataVisualizer Agent. Your role is to:
            1. Create informative and professional visualizations
            2. Generate appropriate charts for different data types
            3. Ensure plots are publication-ready with proper formatting
            4. Provide insights about what each visualization reveals
            
            When creating visualizations, focus on:
            - Distribution plots for numerical variables (histograms, box plots)
            - Correlation analysis with heatmaps
            - Categorical variable distributions (bar charts, pie charts)
            - Time series analysis if applicable
            - Relationship plots between variables
            
            Always create clear, meaningful visualizations that enhance understanding.
            Use the generate_visualizations function to create comprehensive plots."""
        )
    
    def generate_visualizations(self, data: pd.DataFrame) -> List[str]:
        """Generate comprehensive visualizations for the dataset."""
        plots = []
        
        try:
            # Create output directory
            plots_dir = 'eda_plots'
            if not os.path.exists(plots_dir):
                os.makedirs(plots_dir)
            
            # Set style
            plt.style.use('default')
            sns.set_palette("husl")
            
            # 1. Distribution plots for numerical variables
            numerical_cols = data.select_dtypes(include=[np.number]).columns
            if len(numerical_cols) > 0:
                plots.extend(self._create_numerical_plots(data, numerical_cols, plots_dir))
            
            # 2. Categorical variable analysis
            categorical_cols = data.select_dtypes(include=['object']).columns
            if len(categorical_cols) > 0:
                plots.extend(self._create_categorical_plots(data, categorical_cols, plots_dir))
            
            # 3. Correlation analysis
            if len(numerical_cols) > 1:
                plots.extend(self._create_correlation_plots(data, numerical_cols, plots_dir))
            
            # 4. Time series analysis (if applicable)
            time_cols = [col for col in data.columns if 'date' in col.lower() or 'time' in col.lower()]
            if time_cols:
                plots.extend(self._create_time_series_plots(data, time_cols, plots_dir))
            
            print(f"📊 Generated {len(plots)} visualizations in '{plots_dir}' directory")
            
        except Exception as e:
            print(f"⚠️ Warning: Could not create visualizations: {str(e)}")
        
        return plots
    
    def _create_numerical_plots(self, data: pd.DataFrame, numerical_cols: pd.Index, plots_dir: str) -> List[str]:
        """Create distribution plots for numerical variables."""
        plots = []
        
        # Create subplot grid
        n_cols = min(3, len(numerical_cols))
        n_rows = (len(numerical_cols) + n_cols - 1) // n_cols
        
        fig, axes = plt.subplots(n_rows, n_cols, figsize=(15, 5 * n_rows))
        fig.suptitle('Distribution of Numerical Variables', fontsize=16)
        
        if n_rows == 1:
            axes = [axes] if n_cols == 1 else axes
        else:
            axes = axes.flatten()
        
        for i, col in enumerate(numerical_cols):
            if i < len(axes):
                # Histogram with KDE
                sns.histplot(data[col], kde=True, ax=axes[i], bins=30, alpha=0.7)
                axes[i].set_title(f'{col} Distribution')
                axes[i].set_xlabel(col)
                axes[i].set_ylabel('Frequency')
                
                # Add mean line
                mean_val = data[col].mean()
                axes[i].axvline(mean_val, color='red', linestyle='--', 
                               label=f'Mean: {mean_val:.2f}')
                axes[i].legend()
        
        # Hide empty subplots
        for i in range(len(numerical_cols), len(axes)):
            axes[i].set_visible(False)
        
        plt.tight_layout()
        plot_path = f'{plots_dir}/numerical_distributions.png'
        plt.savefig(plot_path, dpi=300, bbox_inches='tight')
        plt.close()
        plots.append(plot_path)
        
        return plots
    
    def _create_categorical_plots(self, data: pd.DataFrame, categorical_cols: pd.Index, plots_dir: str) -> List[str]:
        """Create plots for categorical variables."""
        plots = []
        
        for col in categorical_cols:
            plt.figure(figsize=(10, 6))
            
            # Bar plot
            value_counts = data[col].value_counts().head(15)
            value_counts.plot(kind='bar')
            plt.title(f'{col} Distribution')
            plt.xlabel(col)
            plt.ylabel('Count')
            plt.xticks(rotation=45, ha='right')
            plt.tight_layout()
            
            plot_path = f'{plots_dir}/{col}_distribution.png'
            plt.savefig(plot_path, dpi=300, bbox_inches='tight')
            plt.close()
            plots.append(plot_path)
        
        return plots
    
    def _create_correlation_plots(self, data: pd.DataFrame, numerical_cols: pd.Index, plots_dir: str) -> List[str]:
        """Create correlation analysis plots."""
        plots = []
        
        # Correlation heatmap
        plt.figure(figsize=(10, 8))
        correlation_matrix = data[numerical_cols].corr()
        sns.heatmap(correlation_matrix, annot=True, cmap='coolwarm', center=0, 
                   square=True, linewidths=0.5, fmt='.2f')
        plt.title('Correlation Matrix of Numerical Variables')
        plt.tight_layout()
        
        plot_path = f'{plots_dir}/correlation_heatmap.png'
        plt.savefig(plot_path, dpi=300, bbox_inches='tight')
        plt.close()
        plots.append(plot_path)
        
        return plots
    
    def _create_time_series_plots(self, data: pd.DataFrame, time_cols: List[str], plots_dir: str) -> List[str]:
        """Create time series plots if applicable."""
        plots = []
        
        for col in time_cols:
            try:
                # Convert to datetime if possible
                if data[col].dtype == 'object':
                    data[col] = pd.to_datetime(data[col], errors='coerce')
                
                if data[col].dtype == 'datetime64[ns]':
                    plt.figure(figsize=(12, 6))
                    
                    # Time series plot
                    data[col].value_counts().sort_index().plot(kind='line', marker='o')
                    plt.title(f'{col} - Time Series Analysis')
                    plt.xlabel('Date')
                    plt.ylabel('Frequency')
                    plt.xticks(rotation=45)
                    plt.tight_layout()
                    
                    plot_path = f'{plots_dir}/{col}_timeseries.png'
                    plt.savefig(plot_path, dpi=300, bbox_inches='tight')
                    plt.close()
                    plots.append(plot_path)
                    
            except Exception as e:
                print(f"Could not create time series plot for {col}: {e}")
        
        return plots

class ReportGeneratorAgent(AssistantAgent):
    """Agent responsible for generating comprehensive EDA reports."""
    
    def __init__(self, model_client):
        super().__init__(
            name="ReportGenerator",
            model_client=model_client,
            system_message="""You are the ReportGenerator Agent. Your role is to:
            1. Compile all analysis results into a professional report
            2. Structure the report logically with clear sections
            3. Include all visualizations and statistical findings
            4. Ensure the report is ready for stakeholder presentation
            
            When generating reports, ensure:
            - Executive summary with key findings
            - Comprehensive data overview and quality assessment
            - Detailed statistical analysis with interpretations
            - Visual analysis with plot descriptions
            - Key insights and business implications
            - Clear conclusions and actionable recommendations
            - Professional formatting and structure
            
            Always create comprehensive, well-organized reports that tell a complete data story.
            Use the generate_report function to create the final document."""
        )
    
    def generate_report(self, data: pd.DataFrame, analysis: Dict[str, Any], 
                       plots: List[str], output_path: str = None) -> str:
        """Generate a comprehensive EDA report in .docx format."""
        try:
            print("\n📝 Generating Comprehensive EDA Report...")
            
            # Create document
            doc = Document()
            
            # Title page
            title = doc.add_heading('Exploratory Data Analysis Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            subtitle = doc.add_paragraph('Comprehensive Data Analysis Report')
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            date_para = doc.add_paragraph(f'Generated on: {datetime.now().strftime("%B %d, %Y at %I:%M %p")}')
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_page_break()
            
            # Table of Contents
            toc_heading = doc.add_heading('Table of Contents', level=1)
            toc_content = doc.add_paragraph(
                '1. Executive Summary\n'
                '2. Data Overview\n'
                '3. Data Quality Assessment\n'
                '4. Statistical Analysis\n'
                '5. Visual Analysis\n'
                '6. Key Insights\n'
                '7. Conclusions and Recommendations\n'
                '8. Appendix'
            )
            
            doc.add_page_break()
            
            # 1. Executive Summary
            doc.add_heading('1. Executive Summary', level=1)
            doc.add_paragraph(
                'This report presents a comprehensive exploratory data analysis of the dataset. '
                'The analysis reveals key patterns, trends, and insights that can inform '
                'business decisions and strategic planning. The multi-agent system has '
                'collaboratively analyzed the data to provide actionable insights.'
            )
            
            # 2. Data Overview
            doc.add_heading('2. Data Overview', level=1)
            doc.add_paragraph(f'The dataset contains {data.shape[0]:,} records with {data.shape[1]} variables.')
            
            # Dataset information table
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Variable'
            hdr_cells[1].text = 'Data Type'
            hdr_cells[2].text = 'Description'
            
            for col in data.columns:
                row_cells = table.add_row().cells
                row_cells[0].text = col
                row_cells[1].text = str(data[col].dtype)
                row_cells[2].text = self._get_column_description(col)
            
            doc.add_paragraph()
            
            # 3. Data Quality Assessment
            doc.add_heading('3. Data Quality Assessment', level=1)
            
            # Missing values analysis
            missing_data = data.isnull().sum()
            missing_percent = (missing_data / len(data)) * 100
            
            doc.add_paragraph('Missing Values Analysis:')
            if missing_data.sum() > 0:
                missing_table = doc.add_table(rows=1, cols=3)
                missing_table.style = 'Table Grid'
                missing_hdr = missing_table.rows[0].cells
                missing_hdr[0].text = 'Variable'
                missing_hdr[1].text = 'Missing Count'
                missing_hdr[2].text = 'Missing Percentage'
                
                for col in data.columns:
                    if missing_data[col] > 0:
                        row_cells = missing_table.add_row().cells
                        row_cells[0].text = col
                        row_cells[1].text = str(missing_data[col])
                        row_cells[2].text = f"{missing_percent[col]:.2f}%"
            else:
                doc.add_paragraph('✅ No missing values detected in the dataset.')
            
            doc.add_paragraph()
            
            # 4. Statistical Analysis
            doc.add_heading('4. Statistical Analysis', level=1)
            
            # Numerical variables summary
            numerical_cols = data.select_dtypes(include=[np.number]).columns
            if len(numerical_cols) > 0:
                doc.add_paragraph('Numerical Variables Summary:')
                try:
                    stats_df = data[numerical_cols].describe()
                    
                    # Create stats table
                    if not stats_df.empty and len(stats_df.columns) > 0:
                        stats_table = doc.add_table(rows=1, cols=len(stats_df.index) + 1)
                        stats_table.style = 'Table Grid'
                        stats_hdr = stats_table.rows[0].cells
                        stats_hdr[0].text = 'Statistic'
                        for i, col in enumerate(stats_df.columns):
                            stats_hdr[i + 1].text = str(col)
                        
                        for stat in stats_df.index:
                            row_cells = stats_table.add_row().cells
                            row_cells[0].text = str(stat)
                            for i, col in enumerate(stats_df.columns):
                                try:
                                    value = stats_df.loc[stat, col]
                                    if pd.isna(value):
                                        row_cells[i + 1].text = "N/A"
                                    else:
                                        row_cells[i + 1].text = f"{value:.2f}"
                                except:
                                    row_cells[i + 1].text = "Error"
                        
                        doc.add_paragraph()
                    else:
                        doc.add_paragraph('⚠️ Could not generate numerical statistics table.')
                        
                except Exception as e:
                    doc.add_paragraph(f'⚠️ Error generating numerical statistics: {str(e)}')
                    doc.add_paragraph()
            
            # Categorical variables analysis
            categorical_cols = data.select_dtypes(include=['object']).columns
            if len(categorical_cols) > 0:
                doc.add_paragraph('Categorical Variables Analysis:')
                for col in categorical_cols:
                    try:
                        doc.add_paragraph(f'{col}:')
                        value_counts = data[col].value_counts()
                        
                        if len(value_counts) > 0:
                            cat_table = doc.add_table(rows=1, cols=2)
                            cat_table.style = 'Table Grid'
                            cat_hdr = cat_table.rows[0].cells
                            cat_hdr[0].text = 'Category'
                            cat_hdr[1].text = 'Count'
                            
                            for category, count in value_counts.head(10).items():
                                row_cells = cat_table.add_row().cells
                                row_cells[0].text = str(category) if pd.notna(category) else "Missing"
                                row_cells[1].text = str(count)
                            
                            doc.add_paragraph()
                        else:
                            doc.add_paragraph(f'⚠️ No data available for {col}')
                            
                    except Exception as e:
                        doc.add_paragraph(f'⚠️ Error analyzing {col}: {str(e)}')
                        doc.add_paragraph()
            
            # 5. Visual Analysis
            doc.add_heading('5. Visual Analysis', level=1)
            doc.add_paragraph('The following visualizations provide insights into the data patterns and relationships.')
            
            # Add visualization descriptions
            doc.add_paragraph('Key Visualizations Generated:')
            for plot in plots:
                plot_name = os.path.basename(plot).replace('.png', '').replace('_', ' ').title()
                doc.add_paragraph(f'• {plot_name}', style='List Bullet')
            
            # 6. Key Insights
            doc.add_heading('6. Key Insights', level=1)
            insights = self._generate_insights(data, analysis)
            for insight in insights:
                doc.add_paragraph(f'• {insight}', style='List Bullet')
            
            # 7. Conclusions and Recommendations
            doc.add_heading('7. Conclusions and Recommendations', level=1)
            doc.add_paragraph('Based on the comprehensive analysis, the following conclusions and recommendations are provided:')
            
            conclusions = self._generate_conclusions(data, analysis)
            for conclusion in conclusions:
                doc.add_paragraph(f'• {conclusion}', style='List Bullet')
            
            # 8. Appendix
            doc.add_heading('8. Appendix', level=1)
            doc.add_paragraph('Technical Details and Additional Information')
            
            # Save the document
            if output_path is None:
                output_path = f'EDA_Report_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
            
            doc.save(output_path)
            
            print(f"✅ EDA report generated successfully: {output_path}")
            print(f"📁 Report saved as: {os.path.abspath(output_path)}")
            
            return output_path
            
        except Exception as e:
            error_msg = f"❌ Error generating report: {str(e)}"
            print(error_msg)
            return error_msg
    
    def _get_column_description(self, column_name: str) -> str:
        """Get a description for a given column."""
        descriptions = {
            'customer_id': 'Unique identifier for each customer',
            'age': 'Customer age in years',
            'gender': 'Customer gender (Male/Female)',
            'income': 'Annual income in currency units',
            'purchase_amount': 'Amount spent on purchase',
            'product_category': 'Category of product purchased',
            'customer_satisfaction_score': 'Customer satisfaction rating (1-5 scale)',
            'region': 'Geographic region of customer',
            'purchase_date': 'Date of purchase'
        }
        return descriptions.get(column_name, 'Variable from the dataset')
    
    def _generate_insights(self, data: pd.DataFrame, analysis: Dict[str, Any]) -> List[str]:
        """Generate key insights from the data and analysis."""
        insights = []
        
        try:
            # Basic insights
            insights.append(f"Dataset contains {data.shape[0]:,} records with {data.shape[1]} variables")
            
            # Missing data insights
            missing_data = data.isnull().sum()
            if missing_data.sum() > 0:
                insights.append(f"Missing data detected in {missing_data[missing_data > 0].count()} variables")
            else:
                insights.append("Data quality is excellent with no missing values")
            
            # Numerical insights
            numerical_cols = data.select_dtypes(include=[np.number]).columns
            if len(numerical_cols) > 0:
                insights.append(f"Numerical variables analyzed: {', '.join(numerical_cols)}")
                
                # Income insights
                if 'income' in data.columns:
                    avg_income = data['income'].mean()
                    insights.append(f"Average customer income: ${avg_income:,.2f}")
                
                # Purchase amount insights
                if 'purchase_amount' in data.columns:
                    avg_purchase = data['purchase_amount'].mean()
                    insights.append(f"Average purchase amount: ${avg_purchase:.2f}")
            
            # Categorical insights
            categorical_cols = data.select_dtypes(include=['object']).columns
            if len(categorical_cols) > 0:
                insights.append(f"Categorical variables analyzed: {', '.join(categorical_cols)}")
                
                # Gender distribution
                if 'gender' in data.columns:
                    gender_dist = data['gender'].value_counts()
                    insights.append(f"Gender distribution: {gender_dist.to_dict()}")
                
                # Product category insights
                if 'product_category' in data.columns:
                    top_category = data['product_category'].value_counts().index[0]
                    insights.append(f"Most popular product category: {top_category}")
            
        except Exception as e:
            insights.append(f"Error generating insights: {str(e)}")
        
        return insights
    
    def _generate_conclusions(self, data: pd.DataFrame, analysis: Dict[str, Any]) -> List[str]:
        """Generate conclusions and recommendations."""
        conclusions = []
        
        try:
            conclusions.append("The dataset provides comprehensive information for business analysis")
            
            # Data quality conclusions
            missing_data = data.isnull().sum()
            if missing_data.sum() == 0:
                conclusions.append("Data quality is excellent with no missing values")
            else:
                conclusions.append(f"Data quality assessment completed with {missing_data.sum()} missing values identified")
            
            # Business insights
            if 'income' in data.columns and 'purchase_amount' in data.columns:
                conclusions.append("Customer income and purchase behavior analysis reveals spending patterns")
            
            if 'customer_satisfaction_score' in data.columns:
                conclusions.append("Customer satisfaction metrics provide insights into service quality")
            
            conclusions.append("Recommendations include further segmentation analysis and predictive modeling")
            conclusions.append("Regular data updates and quality monitoring are recommended")
            
        except Exception as e:
            conclusions.append(f"Error generating conclusions: {str(e)}")
        
        return conclusions

class MultiAgentEDASystem:
    """Main system that orchestrates the multi-agent EDA workflow using AutoGen AgentChat."""
    
    def __init__(self, model_name: str = "gpt-4o", api_key: str = None):
        """Initialize the multi-agent EDA system."""
        
        # Initialize the model client
        self.model_client = OpenAIChatCompletionClient(
            model=model_name,
            api_key=api_key
        )
        
        # Initialize all agents
        self.agents = self._create_agents()
        
        # Create user proxy
        self.user_proxy = UserProxyAgent(name="user_proxy")
        
        # Set up group chat
        self.setup_group_chat()
    
    def _create_agents(self) -> Dict[str, AssistantAgent]:
        """Create all specialized agents for the EDA workflow."""
        agents = {}
        
        # Create specialized agents with model client
        agents['data_loader'] = DataLoaderAgent(self.model_client)
        agents['data_analyzer'] = DataAnalyzerAgent(self.model_client)
        agents['data_visualizer'] = DataVisualizerAgent(self.model_client)
        agents['report_generator'] = ReportGeneratorAgent(self.model_client)
        
        return agents
    
    def setup_group_chat(self):
        """Set up the group chat for agent collaboration."""
        # Combine all agents
        all_agents = [self.user_proxy] + list(self.agents.values())
        
        # Create group chat with RoundRobinGroupChat
        self.group_chat = RoundRobinGroupChat(
            all_agents,
            termination_condition=TextMentionTermination("TERMINATE_EDA")
        )
    
    async def run_eda_workflow(self, file_path: str) -> Dict[str, Any]:
        """Execute the complete EDA workflow using AutoGen AgentChat."""
        try:
            print("🚀 Starting Multi-Agent EDA System (AutoGen AgentChat)...")
            print("=" * 60)
            
            # Check if file exists
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"File not found at {file_path}")
            
            print(f"✅ Dataset file found: {file_path}")
            
            # Step 1: Load and validate data
            print("\n📊 Step 1: Loading and validating data...")
            data = self.agents['data_loader'].load_data(file_path)
            data_info = self.agents['data_loader'].get_data_info(data)
            
            print(f"   - Shape: {data.shape}")
            print(f"   - Columns: {', '.join(data.columns.tolist())}")
            print(f"   - Memory usage: {data_info['memory_usage']:.2f} KB")
            
            # Step 2: Perform data analysis
            print("\n🔍 Step 2: Performing comprehensive data analysis...")
            analysis = self.agents['data_analyzer'].analyze_data(data)
            
            print("   - Statistical analysis completed")
            print("   - Data quality assessment finished")
            
            # Step 3: Generate visualizations
            print("\n📈 Step 3: Creating data visualizations...")
            plots = self.agents['data_visualizer'].generate_visualizations(data)
            
            print(f"   - Generated {len(plots)} visualizations")
            
            # Step 4: Generate comprehensive report
            print("\n📝 Step 4: Generating comprehensive EDA report...")
            report_path = self.agents['report_generator'].generate_report(
                data, analysis, plots
            )
            
            # Return results
            results = {
                "data": data,
                "analysis": analysis,
                "plots": plots,
                "report_path": report_path,
                "status": "success"
            }
            
            print("\n" + "=" * 60)
            print("🎉 Multi-Agent EDA Process Complete!")
            print(f"📁 Final Report: {report_path}")
            
            return results
            
        except Exception as e:
            error_msg = f"❌ Error in EDA workflow: {str(e)}"
            print(error_msg)
            return {"status": "error", "message": error_msg}
    
    async def run_eda_workflow_dataframe(self, data: pd.DataFrame) -> Dict[str, Any]:
        """Execute the complete EDA workflow directly with a DataFrame."""
        try:
            print("🚀 Starting Multi-Agent EDA System (AutoGen AgentChat) with DataFrame...")
            print("=" * 60)
            
            print(f"✅ Dataset loaded directly: {data.shape[0]} rows × {data.shape[1]} columns")
            
            # Step 1: Validate and clean data
            print("\n📊 Step 1: Validating and cleaning data...")
            
            # Handle duplicate column names if they exist
            if len(data.columns) != len(set(data.columns)):
                print("⚠️ Warning: Duplicate column names detected. Renaming duplicates...")
                duplicate_cols = data.columns[data.columns.duplicated()].tolist()
                print(f"   Duplicate columns: {duplicate_cols}")
                
                new_columns = []
                seen_columns = {}
                
                for col in data.columns:
                    if col in seen_columns:
                        seen_columns[col] += 1
                        new_columns.append(f"{col}.{seen_columns[col]}")
                    else:
                        seen_columns[col] = 0
                        new_columns.append(col)
                
                data.columns = new_columns
                print(f"   Columns renamed to: {list(data.columns)}")
            
            # Clean empty rows and columns
            initial_rows, initial_cols = data.shape
            data = data.dropna(how='all')
            data = data.dropna(axis=1, how='all')
            print(f"   Cleaned data shape: {data.shape[0]} rows × {data.shape[1]} columns")
            
            data_info = self.agents['data_loader'].get_data_info(data)
            print(f"   - Memory usage: {data_info['memory_usage']:.2f} KB")
            
            # Step 2: Perform data analysis
            print("\n🔍 Step 2: Performing comprehensive data analysis...")
            analysis = self.agents['data_analyzer'].analyze_data(data)
            
            print("   - Statistical analysis completed")
            print("   - Data quality assessment finished")
            
            # Step 3: Generate visualizations
            print("\n📈 Step 3: Creating data visualizations...")
            plots = self.agents['data_visualizer'].generate_visualizations(data)
            
            print(f"   - Generated {len(plots)} visualizations")
            
            # Step 4: Generate comprehensive report
            print("\n📝 Step 4: Generating comprehensive EDA report...")
            report_path = self.agents['report_generator'].generate_report(
                data, analysis, plots
            )
            
            # Return results
            results = {
                "data": data,
                "analysis": analysis,
                "plots": plots,
                "report_path": report_path,
                "status": "success"
            }
            
            print("\n" + "=" * 60)
            print("🎉 Multi-Agent EDA Process Complete!")
            print(f"📁 Final Report: {report_path}")
            
            return results
            
        except Exception as e:
            error_msg = f"❌ Error in EDA workflow: {str(e)}"
            print(error_msg)
            return {"status": "error", "message": error_msg}
    
    async def cleanup(self):
        """Clean up resources."""
        await self.model_client.close()

def get_csv_file_path():
    """Interactive function to get CSV file path from user."""
    print("\n" + "="*60)
    print("📁 CSV Dataset Selection")
    print("="*60)
    
    while True:
        print("\nOptions:")
        print("1. Use default sample dataset")
        print("2. Upload CSV file to current directory")
        print("3. Enter custom file path")
        print("4. Exit")
        
        choice = input("\nSelect option (1-4): ").strip()
        
        if choice == "1":
            # Default sample dataset
            default_path = r"C:\Users\mohsi\aaai\EDA_MAS\sample_customer_data.csv"
            if os.path.exists(default_path):
                print(f"✅ Using default dataset: {default_path}")
                return default_path
            else:
                print(f"❌ Default dataset not found at: {default_path}")
                print("Please use option 2 or 3 instead.")
                continue
                
        elif choice == "2":
            # File upload to current directory
            print("\n📤 File Upload Instructions:")
            print("1. Place your CSV file in the current directory")
            print("2. Enter the filename (e.g., 'my_data.csv')")
            print("3. Or drag and drop your CSV file here")
            
            filename = input("\nEnter CSV filename: ").strip().strip('"')  # Remove quotes if user drags file
            
            # Check if file exists
            if os.path.exists(filename):
                if filename.lower().endswith('.csv'):
                    print(f"✅ File found: {filename}")
                    return filename
                else:
                    print("❌ File must have .csv extension")
                    continue
            else:
                print(f"❌ File not found: {filename}")
                print("Please check the filename and try again.")
                continue
                
        elif choice == "3":
            # Custom file path
            file_path = input("Enter full file path: ").strip().strip('"')
            
            if os.path.exists(file_path):
                if file_path.lower().endswith('.csv'):
                    print(f"✅ File found: {file_path}")
                    return file_path
                else:
                    print("❌ File must have .csv extension")
                    continue
            else:
                print(f"❌ File not found: {file_path}")
                print("Please check the file path and try again.")
                continue
                
        elif choice == "4":
            print("👋 Goodbye!")
            return None
            
        else:
            print("❌ Invalid choice. Please select 1, 2, 3, or 4.")
            continue

# Main execution function
async def main():
    """Main function to run the EDA system."""
    
    print("🚀 Multi-Agent EDA System")
    print("=" * 50)
    
    # Get CSV file path using the helper function
    file_path = get_csv_file_path()
    
    if file_path is None:
        print("👋 Exiting...")
        return
    
    print(f"\n📊 Analyzing dataset: {file_path}")
    
    # Initialize the EDA system
    eda_system = MultiAgentEDASystem()
    
    try:
        # Run the EDA workflow
        results = await eda_system.run_eda_workflow(file_path)
        
        if results["status"] == "success":
            print(f"\n✅ EDA completed successfully!")
            print(f"📊 Dataset analyzed: {results['data'].shape}")
            print(f"📈 Visualizations created: {len(results['plots'])}")
            print(f"📝 Report generated: {results['report_path']}")
            
            # Show where files are saved
            print(f"\n📁 Files generated:")
            print(f"   - EDA Report: {results['report_path']}")
            print(f"   - Visualizations: eda_plots/ directory")
            
        else:
            print(f"\n❌ EDA failed: {results['message']}")
    
    finally:
        # Clean up
        await eda_system.cleanup()

# Run the system
if __name__ == "__main__":
    # Note: You'll need to set your OpenAI API key
    # export OPENAI_API_KEY="your-api-key-here"
    # Or create a .env file with: OPENAI_API_KEY=your-api-key-here
    asyncio.run(main())
